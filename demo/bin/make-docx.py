#!/usr/bin/env -S uv run --script
# /// script
# dependencies = ["python-docx"]
# ///
"""Word document builder CLI connector for Cortex Cowork.

Invoked by the cowork-runner as a CLI connector tool: the agent writes a JSON
spec into the sandbox, then calls this script to render it as .docx:

  make-docx.py work/spec.json --out artifacts/Uzasadnienia_kosztow.docx

Spec format:

  {
    "title": "Document title",
    "blocks": [
      {"type": "heading", "text": "Zadanie 1 - ..."},
      {"type": "subheading", "text": "Stanowisko | 50% etatu"},
      {"type": "paragraph", "text": "..."},
      {"type": "bullets", "items": ["...", "..."]},
      {"type": "field", "label": "Uzasadnienie kosztu", "limit": 1500, "text": "..."},
      {"type": "warning", "text": "Suma etatów: 125% - sprawdź..."},
      {"type": "rule"}
    ]
  }

"field" renders "<label> (limit N znaków):" plus the text (lines starting
with "- " become a bullet list) plus an italic counter line "[X znaków z N]".
Fields are validated BEFORE the file is written: length is measured on
whitespace-collapsed text, and specs still containing "[WPISZ" placeholders
get a +150 char allowance for what the user will type in. Any field
projected over its limit aborts with exit code 2 and a per-field status
report (green/yellow/red), so the agent can shorten and retry.

Guardrail: when COWORK_SANDBOX_DIR is present in the environment (it always
is when spawned by the runner), both the spec file and --out must resolve
inside that sandbox - a connector tool must not read or write anywhere else
on the host.
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

PLACEHOLDER_OVERHEAD = 150
PLACEHOLDER_MARK = "[WPISZ"


def fail(message: str, code: int = 1) -> "sys.NoReturn":
    print(f"error: {message}", file=sys.stderr)
    sys.exit(code)


def guard_sandbox(path: Path, role: str) -> Path:
    resolved = path.resolve()
    sandbox = os.environ.get("COWORK_SANDBOX_DIR")
    if sandbox:
        sandbox_root = Path(sandbox).resolve()
        if not resolved.is_relative_to(sandbox_root):
            fail(f"{role} must point inside the session sandbox ({sandbox_root}); got {resolved}")
    return resolved


def collapsed_length(text: str) -> int:
    return len(re.sub(r"\s+", " ", text).strip())


def check_field(label: str, text: str, limit: int) -> dict:
    length = collapsed_length(text)
    projected = length + PLACEHOLDER_OVERHEAD if PLACEHOLDER_MARK in text else length
    if projected > limit:
        status = f"🔴 PRZEKROCZY limit po uzupełnieniu ({projected} szac. / {limit})"
    elif projected > limit * 0.9:
        status = f"🟡 Blisko limitu po uzupełnieniu ({projected} szac. / {limit})"
    else:
        status = "🟢 OK"
    return {"label": label, "len": length, "projected": projected, "ok": projected <= limit, "status": status}


def add_rule(document) -> None:
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    border = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_pr.makeelement(
        qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "999999"},
    )
    border.append(bottom)
    p_pr.append(border)


def add_field_text(document, text: str) -> None:
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(("- ", "– ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)


def render(spec: dict, checks: list[dict], out: Path) -> None:
    import docx

    document = docx.Document()
    if spec.get("title"):
        document.add_heading(str(spec["title"]), level=0)

    check_iter = iter(checks)
    for block in spec.get("blocks", []):
        kind = block.get("type")
        if kind == "heading":
            document.add_heading(str(block.get("text", "")), level=1)
        elif kind == "subheading":
            document.add_heading(str(block.get("text", "")), level=2)
            add_rule(document)
        elif kind == "paragraph":
            document.add_paragraph(str(block.get("text", "")))
        elif kind == "bullets":
            for item in block.get("items", []):
                document.add_paragraph(str(item), style="List Bullet")
        elif kind == "field":
            check = next(check_iter)
            limit = int(block.get("limit", 1500))
            label_paragraph = document.add_paragraph()
            label_paragraph.add_run(f"{block.get('label', 'Pole')} (limit {limit} znaków):").bold = True
            add_field_text(document, str(block.get("text", "")))
            counter = document.add_paragraph()
            run = counter.add_run(f"[{check['len']} znaków z {limit}]")
            run.italic = True
        elif kind == "warning":
            paragraph = document.add_paragraph()
            paragraph.add_run(f"⚠️ {block.get('text', '')}").bold = True
        elif kind == "rule":
            add_rule(document)
        else:
            fail(f"unknown block type: {kind!r}")

    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out))


def main() -> None:
    parser = argparse.ArgumentParser(description="Render a JSON spec as a .docx in the session sandbox")
    parser.add_argument("spec", help="Path to the JSON spec (inside the session sandbox)")
    parser.add_argument("--out", required=True, help="Output .docx path (inside the session sandbox)")
    args = parser.parse_args()

    spec_path = guard_sandbox(Path(args.spec), "spec")
    out = guard_sandbox(Path(args.out), "--out")
    if not spec_path.is_file():
        fail(f"no such file: {spec_path}")

    try:
        spec = json.loads(spec_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as error:
        fail(f"spec is not valid JSON: {error}")

    fields = [block for block in spec.get("blocks", []) if block.get("type") == "field"]
    checks = [
        check_field(str(block.get("label", "Pole")), str(block.get("text", "")), int(block.get("limit", 1500)))
        for block in fields
    ]

    if checks:
        print("=== Walidacja długości pól ===")
        for check in checks:
            print(f"  {check['label']}: {check['len']} znaków - {check['status']}")

    failed = [check for check in checks if not check["ok"]]
    if failed:
        print(f"\n⚠️  {len(failed)} pole(a) prawdopodobnie przekroczy limit po uzupełnieniu:", file=sys.stderr)
        for check in failed:
            print(f"   - {check['label']}", file=sys.stderr)
        fail("skróć wskazane pola w spec i wywołaj ponownie - plik NIE został wygenerowany", code=2)

    render(spec, checks, out)
    print(f"✅ Zapisano {args.out}" + (f" ({len(checks)} pól zwalidowanych)" if checks else ""))


if __name__ == "__main__":
    main()
